from flask import Flask, request, render_template, send_file
from docx import Document
from docx.shared import Inches
import os

app = Flask(__name__)
UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'outputs'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# Function to parse the .docx file and extract questions
def parse_docx(file_path):
    doc = Document(file_path)
    questions = []
    question_data = None

    option_prefixes = [" ", "a.", "b.", "c.", "d.", "a)", "b)", "c)", "d)", "A.", "B.", "C.", "D.", "A)", "B)", "C)", "D)", "(a)", "(b)", "(c)", "(d)", "(A)", "(B)", "(C)", "(D)"]
    opt_pref = ["(a)", "(b)", "(c)", "(d)"]
    i = 0

    for para in doc.paragraphs:
        text = para.text.strip()

        answer_prefixes = ["Answer: ", "उत्तर ", "Answer-","Answer"]

        for prefix in answer_prefixes:
            if prefix in text:
                if question_data:
                    question_data["answer"] = text.split(prefix, 1)[-1].strip()
                break
        else:
            if any(text.startswith(prefix) for prefix in option_prefixes):
                if question_data:
                    option_key = next(prefix for prefix in option_prefixes if text.startswith(prefix))
                    question_data["options"][option_key] = text.split(option_key, 1)[-1].strip()

            elif text and not question_data:
                if text[0].isdigit():
                    question_data = {
                        "question": text,
                        "options": {},
                        "answer": None,
                        "solution": "",
                        "positive_marks": "2",
                        "negative_marks": "0"
                    }
                else:
                    question_data = {
                        "question": text,
                        "options": {},
                        "answer": None,
                        "solution": "",
                        "positive_marks": "2",
                        "negative_marks": "0"
                    }
            elif text and question_data:
                if "Explanation:" in text or "व्याख्या" in text:
                    try:
                        question_data["solution"] = text.split(":", 1)[1].strip()
                    except IndexError:
                        print(f"Warning: Unable to split text to extract solution: {text}")
                elif len(question_data["options"]) < 4:
                    option_key = opt_pref[i % 4]
                    i += 1
                    question_data["options"][option_key] = text.strip()
                else:
                    question_data["solution"] += f"{text} "

            if not text and question_data and question_data["solution"]:
                questions.append(question_data)
                question_data = None

    if question_data:
        questions.append(question_data)

    return questions

# Helper function to get options
def get_option(options, keys):
    for key in keys:
        if key in options:
            return options[key]
    return ""

def set_col_widths(table):
    # Adjust widths: 30% for the field column, 70% for the value column
    widths = [Inches(1.8), Inches(4.2)]
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]

# Function to generate a .docx file with a table for each question
def generate_docx(questions, output_path):
    doc = Document()

    for q in questions:
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'

        set_col_widths(table)

        fields = [
            ("Question", q['question']),
            ("Type", "multiple_choice"),
            ("Option", get_option(q["options"], [" ", "a.", "a)", "A.", "A)", "(a)", "(A)"])),
            ("Option", get_option(q["options"], [" ", "b.", "b)", "B.", "B)", "(b)", "(B)"])),
            ("Option", get_option(q["options"], [" ", "c.", "c)", "C.", "C)", "(c)", "(C)"])),
            ("Option", get_option(q["options"], [" ", "d.", "d)", "D.", "D)", "(d)", "(D)"])),
            ("Answer", q['answer']),
            ("Solution", q['solution']),
            ("Positive Marks", q['positive_marks']),
            ("Negative Marks", q['negative_marks']),
        ]

        for field, value in fields:
            row_cells = table.add_row().cells
            row_cells[0].text = field
            row_cells[1].text = "" if value is None else str(value)

        doc.add_paragraph("\n")

    doc.save(output_path)

# Route for file upload and processing
@app.route('/', methods=['GET', 'POST'])
def upload_file():
    if request.method == 'POST':
        if 'file' not in request.files:
            return "No file part"
        file = request.files['file']
        if file.filename == '':
            return "No selected file"
        if file:
            input_file = os.path.join(UPLOAD_FOLDER, file.filename)
            file.save(input_file)

            try:
                output_file = os.path.join(OUTPUT_FOLDER, f"output_{file.filename}")
                questions = parse_docx(input_file)
                generate_docx(questions, output_file)
                return send_file(output_file, as_attachment=True)
            except Exception as e:
                print(f"Error processing file: {e}")
                return "An error occurred while processing the file."

    return render_template('upload.html')

if __name__ == '__main__':
    app.run(debug=True)
